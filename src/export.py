"""Export functionality for PDF and DOCX formats."""

import io
from typing import Dict, Any, Optional
from datetime import datetime
from enum import Enum


class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"


class DocumentExporter:
    """Handles document export to various formats."""
    
    def __init__(self):
        """Initialize the document exporter."""
        self._validate_dependencies()
    
    def _validate_dependencies(self):
        """Validate that required libraries are available."""
        try:
            import reportlab
            import docx
        except ImportError as e:
            raise ImportError(
                "Missing required dependencies for export. "
                "Install with: pip install reportlab python-docx"
            ) from e
    
    def export_to_pdf(
        self,
        content: Dict[str, Any],
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export document content to PDF format.
        
        Args:
            content: Document content with structure and text
            metadata: Optional metadata (title, author, etc.)
            
        Returns:
            PDF file as bytes
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        story = []
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Justify',
            alignment=TA_JUSTIFY,
            fontSize=12,
            leading=14
        ))
        
        if metadata:
            title = metadata.get('title', 'Untitled Document')
            author = metadata.get('author', 'Unknown')
            date = metadata.get('date', datetime.now().strftime('%Y-%m-%d'))
            
            title_style = ParagraphStyle(
                name='CustomTitle',
                fontSize=24,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            story.append(Paragraph(title, title_style))
            story.append(Paragraph(f"by {author}", styles['Normal']))
            story.append(Paragraph(date, styles['Normal']))
            story.append(Spacer(1, 0.5 * inch))
        
        text_content = content.get('text', '')
        chapters = content.get('chapters', [])
        
        if chapters:
            for chapter in chapters:
                chapter_title = chapter.get('title', '')
                chapter_text = chapter.get('text', '')
                
                if chapter_title:
                    story.append(Paragraph(chapter_title, styles['Heading1']))
                    story.append(Spacer(1, 0.2 * inch))
                
                paragraphs = chapter_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para, styles['Justify']))
                        story.append(Spacer(1, 0.1 * inch))
                
                story.append(Spacer(1, 0.3 * inch))
        else:
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para, styles['Justify']))
                    story.append(Spacer(1, 0.1 * inch))
        
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
    
    def export_to_docx(
        self,
        content: Dict[str, Any],
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export document content to DOCX format.
        
        Args:
            content: Document content with structure and text
            metadata: Optional metadata (title, author, etc.)
            
        Returns:
            DOCX file as bytes
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        if metadata:
            title = metadata.get('title', 'Untitled Document')
            author = metadata.get('author', 'Unknown')
            date = metadata.get('date', datetime.now().strftime('%Y-%m-%d'))
            
            title_para = doc.add_paragraph(title)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(24)
            title_run.bold = True
            
            author_para = doc.add_paragraph(f"by {author}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(date)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
        
        text_content = content.get('text', '')
        chapters = content.get('chapters', [])
        
        if chapters:
            for chapter in chapters:
                chapter_title = chapter.get('title', '')
                chapter_text = chapter.get('text', '')
                
                if chapter_title:
                    heading = doc.add_heading(chapter_title, level=1)
                
                paragraphs = chapter_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                
                doc.add_paragraph()
        else:
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
        
        core_properties = doc.core_properties
        if metadata:
            core_properties.title = metadata.get('title', '')
            core_properties.author = metadata.get('author', '')
            core_properties.created = datetime.now()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_data = buffer.getvalue()
        buffer.close()
        
        return docx_data
    
    def export(
        self,
        content: Dict[str, Any],
        format: ExportFormat,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export document to specified format.
        
        Args:
            content: Document content
            format: Export format (PDF or DOCX)
            metadata: Optional document metadata
            
        Returns:
            Exported file as bytes
            
        Raises:
            ValueError: If format is not supported
        """
        if format == ExportFormat.PDF:
            return self.export_to_pdf(content, metadata)
        elif format == ExportFormat.DOCX:
            return self.export_to_docx(content, metadata)
        else:
            raise ValueError(f"Unsupported export format: {format}")
